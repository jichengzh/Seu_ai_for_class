# -*- coding: utf-8 -*-
"""
TaskParser — 解析参考任务书文档，提取章节结构。

支持 .docx 和 .pdf 格式。返回标准化的结构字典，供 TaskGenerator 使用。
"""

from pathlib import Path
from typing import Any


SECTION_PATTERNS: dict[str, list[str]] = {
    "cover":        ["课程名称", "专业", "实习时间", "封面", "课程名"],
    "objectives":   ["课程目标", "设计目标", "学习目标", "课程背景", "背景与目标"],
    "modules":      ["模块", "阶段", "任务概述", "模块概述"],
    "details":      ["设计内容", "任务内容", "具体要求", "各模块"],
    "requirements": ["作业要求", "报告要求", "提交格式", "作业格式"],
    "deliverables": ["提交成果", "交付物", "成果要求"],
    "grading":      ["成绩考核", "评分标准", "考核方式", "成绩评定"],
    "schedule":     ["时间安排", "进度计划", "时间计划"],
    "references":   ["参考资料", "教材", "附录", "参考文献", "参考资源"],
}

MAX_SUMMARY_LEN = 2000


def _match_section(text: str) -> str | None:
    """Return the section key if text matches any pattern, else None."""
    text_lower = text.lower()
    for key, patterns in SECTION_PATTERNS.items():
        for p in patterns:
            if p in text:
                return key
    return None


class TaskParser:
    """解析参考任务书，提取文档结构。"""

    def parse_docx(self, file_path: str) -> dict[str, Any]:
        """
        解析 .docx 文件，返回结构化内容。

        Returns:
            {
                "sections": {"cover": {"title": ..., "content": ...}, ...},
                "module_count": int,
                "has_grading_table": bool,
                "raw_text": str,
            }

        Raises:
            FileNotFoundError: 文件不存在
            Exception: 文件损坏或格式不符
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        from docx import Document  # python-docx

        doc = Document(str(path))  # raises if not a valid docx

        sections: dict[str, dict[str, str]] = {}
        raw_lines: list[str] = []
        current_section: str | None = None
        current_title: str = ""
        current_content_lines: list[str] = []
        has_grading_table = False

        def _flush():
            nonlocal current_section, current_content_lines, current_title
            if current_section:
                sections[current_section] = {
                    "title": current_title,
                    "content": "\n".join(current_content_lines).strip(),
                }
            current_section = None
            current_title = ""
            current_content_lines = []

        # Iterate paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                if current_content_lines:
                    current_content_lines.append("")
                continue

            raw_lines.append(text)

            # Check if this paragraph is a section heading
            is_heading = para.style.name.startswith("Heading") if para.style else False
            matched = _match_section(text)

            if matched and (is_heading or len(text) < 40):
                _flush()
                current_section = matched
                current_title = text
            elif current_section:
                current_content_lines.append(text)

        _flush()

        # Check tables for grading table
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join(c.text for c in row.cells)
                raw_lines.append(row_text)
                if any(kw in row_text for kw in ["优秀", "良好", "分", "评分"]):
                    has_grading_table = True
                # Try to match table content to sections
                matched = _match_section(row_text)
                if matched and matched not in sections:
                    sections[matched] = {"title": row_text[:50], "content": row_text}

        # Estimate module count from modules section content
        module_count = 0
        if "modules" in sections:
            content = sections["modules"]["content"]
            import re
            module_count = len(re.findall(r"模块[一二三四五六七八九十\d]", content))
            if module_count == 0:
                # count numbered items
                module_count = len(re.findall(r"^[1-9][.、]", content, re.MULTILINE))

        return {
            "sections": sections,
            "module_count": max(module_count, 0),
            "has_grading_table": has_grading_table,
            "raw_text": "\n".join(raw_lines),
        }

    def parse_pdf(self, file_path: str) -> dict[str, Any]:
        """
        解析 .pdf 文件（使用 PyMuPDF）。

        Returns:
            同 parse_docx，scanned PDF 时返回 {"error": "scanned_pdf", "raw_text": "", "sections": {}}
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            import fitz  # PyMuPDF
        except ImportError as e:
            raise ImportError("PyMuPDF (fitz) is required for PDF parsing: pip install pymupdf") from e

        doc = fitz.open(str(path))
        raw_lines: list[str] = []

        for page in doc:
            text = page.get_text()
            if text:
                raw_lines.extend(text.splitlines())

        doc.close()

        if not any(line.strip() for line in raw_lines):
            return {"error": "scanned_pdf", "raw_text": "", "sections": {}, "module_count": 0, "has_grading_table": False}

        # Re-use text-based section matching
        sections: dict[str, dict[str, str]] = {}
        current_section: str | None = None
        current_title: str = ""
        current_content_lines: list[str] = []

        def _flush():
            nonlocal current_section, current_content_lines, current_title
            if current_section:
                sections[current_section] = {
                    "title": current_title,
                    "content": "\n".join(current_content_lines).strip(),
                }
            current_section = None
            current_title = ""
            current_content_lines = []

        for line in raw_lines:
            text = line.strip()
            if not text:
                continue
            matched = _match_section(text)
            if matched and len(text) < 60:
                _flush()
                current_section = matched
                current_title = text
            elif current_section:
                current_content_lines.append(text)

        _flush()

        return {
            "sections": sections,
            "module_count": 0,
            "has_grading_table": False,
            "raw_text": "\n".join(l.strip() for l in raw_lines if l.strip()),
        }

    def extract_structure_summary(self, parsed: dict[str, Any]) -> str:
        """
        生成结构摘要字符串，用于提示词中描述参考文档的章节结构。
        截断到 MAX_SUMMARY_LEN 字符。
        """
        if parsed.get("error"):
            return f"[解析错误: {parsed['error']}]"

        sections = parsed.get("sections", {})
        lines: list[str] = []
        for key, section in sections.items():
            title = section.get("title", key)
            content_preview = section.get("content", "")[:200]
            lines.append(f"### {title}\n{content_preview}")

        lines.append(f"\n[共 {len(sections)} 个章节，模块数: {parsed.get('module_count', 0)}]")
        summary = "\n\n".join(lines)

        if len(summary) > MAX_SUMMARY_LEN:
            summary = summary[:MAX_SUMMARY_LEN] + "\n...[截断]"

        return summary


__all__ = ["TaskParser", "SECTION_PATTERNS"]
