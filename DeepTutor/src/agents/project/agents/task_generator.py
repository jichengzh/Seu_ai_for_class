# -*- coding: utf-8 -*-
"""
TaskGenerator — 逐章节流式生成新任务书。

流程：
1. 可选 RAG 检索相关知识
2. 可选 Web 搜索获取最新资料
3. 按固定顺序逐章节流式生成（9 个章节）
4. 导出 Markdown + docx
"""

import asyncio
import time
from pathlib import Path
from typing import Any, Callable

import yaml

from src.logging.logger import get_logger
from src.services.llm import factory as llm_factory

logger = get_logger("TaskGenerator")

# 章节顺序与标题（中文）
SECTION_ORDER_ZH = [
    ("cover",        "封面信息"),
    ("objectives",   "课程背景与目标"),
    ("modules",      "模块概述"),
    ("details",      "各模块设计内容"),
    ("requirements", "作业要求"),
    ("deliverables", "提交成果"),
    ("grading",      "成绩考核"),
    ("schedule",     "时间安排"),
    ("references",   "参考资源"),
]

SECTION_ORDER_EN = [
    ("cover",        "Cover Information"),
    ("objectives",   "Background and Objectives"),
    ("modules",      "Module Overview"),
    ("details",      "Module Design Details"),
    ("requirements", "Assignment Requirements"),
    ("deliverables", "Deliverables"),
    ("grading",      "Grading Criteria"),
    ("schedule",     "Schedule"),
    ("references",   "References"),
]

_PROMPTS_DIR = Path(__file__).resolve().parents[1] / "prompts"


def _load_prompts(language: str) -> dict[str, Any]:
    lang = "zh" if language.startswith("zh") else "en"
    prompt_file = _PROMPTS_DIR / lang / "task_generation.yaml"
    with open(prompt_file, encoding="utf-8") as f:
        return yaml.safe_load(f)


class TaskGenerator:
    """
    逐章节流式生成新任务书。

    Args:
        output_dir: 生成文件存储目录
        language: 提示词语言（"zh" 或 "en"）
    """

    def __init__(self, output_dir: str, language: str = "zh"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.language = language
        self._prompts = _load_prompts(language)
        self._section_order = SECTION_ORDER_ZH if language.startswith("zh") else SECTION_ORDER_EN

    async def generate(
        self,
        theme: str,
        reference_structure: dict[str, Any],
        kb_name: str | None,
        web_search: bool,
        ws_callback: Callable,
    ) -> dict[str, Any]:
        """
        生成完整任务书文档。

        Returns:
            {"content": full_md, "md_path": str, "docx_path": str}
        """
        # Step 1: RAG 检索
        if kb_name:
            await ws_callback({"type": "status", "content": f"正在检索知识库 [{kb_name}]..."})
            rag_context = await self._rag_retrieve(theme, kb_name)
        else:
            await ws_callback({"type": "status", "content": "未选择知识库，无外部知识源，将仅凭 LLM 知识生成"})
            rag_context = ""

        # Step 2: Web 搜索
        web_context = ""
        if web_search:
            await ws_callback({"type": "status", "content": "正在进行网络搜索..."})
            web_context = await self._web_search(theme)

        # Step 3: 逐章节流式生成
        sections: dict[str, str] = {}
        system_prompt = self._prompts.get("system", "You are a helpful assistant.")
        section_prompts = self._prompts.get("section_prompts", {})

        for section_key, section_title in self._section_order:
            await ws_callback({"type": "status", "content": f"正在生成章节：{section_title}"})

            # 获取该章节在参考文档中的内容
            ref_sections = reference_structure.get("sections", {})
            ref_section = ref_sections.get(section_key, {})
            reference_content = ref_section.get("content", "") if ref_section else ""

            # 构建 prompt
            prompt_template = section_prompts.get(section_key, "请为主题 {theme} 生成 {section_title} 章节内容。")
            prompt = prompt_template.format(
                theme=theme,
                rag_context=rag_context[:1500] if rag_context else "（无）",
                web_context=web_context[:1000] if web_context else "（无）",
                reference_content=reference_content[:800] if reference_content else "（无参考内容）",
                modules_content=sections.get("modules", "")[:500],
            )

            # 流式生成
            section_content = ""
            try:
                async for chunk in llm_factory.stream(
                    prompt=prompt,
                    system_prompt=system_prompt,
                    temperature=0.7,
                ):
                    section_content += chunk
                    await ws_callback({"type": "chunk", "content": chunk})
            except Exception as e:
                logger.warning(f"Section '{section_key}' generation error: {e}")
                # 保留空内容，继续下一章节
                await ws_callback({"type": "log", "content": f"章节 {section_title} 生成出错: {e}"})

            sections[section_key] = section_content
            await ws_callback({
                "type": "section",
                "section": section_key,
                "content": section_content,
            })

        # Step 4: 组合并导出
        full_md = self._assemble_markdown(theme, sections)
        md_path = self._save_markdown(full_md)
        docx_path = self._export_docx(full_md)

        return {
            "content": full_md,
            "md_path": str(md_path),
            "docx_path": str(docx_path),
        }

    # ------------------------------------------------------------------
    # Private helpers
    # ------------------------------------------------------------------

    async def _rag_retrieve(self, theme: str, kb_name: str) -> str:
        """调用 RAGService 检索与主题相关的知识。"""
        try:
            from src.services.rag.service import RAGService

            service = RAGService()
            result = await service.search(
                query=theme,
                kb_name=kb_name,
                mode="hybrid",
            )
            return result.get("answer") or result.get("content") or ""
        except Exception as e:
            logger.warning(f"RAG retrieval failed for kb '{kb_name}': {e}")
            return ""

    async def _web_search(self, theme: str) -> str:
        """调用 web_search 服务（同步函数，在线程池中运行）。"""
        try:
            from src.services.search import web_search

            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(None, lambda: web_search(theme))
            return result.get("answer", "")
        except Exception as e:
            logger.warning(f"Web search failed: {e}")
            return ""

    def _assemble_markdown(self, theme: str, sections: dict[str, str]) -> str:
        """将各章节内容组合为完整 Markdown 文档。"""
        header = f"# {theme} — 实习任务书\n\n"
        parts = [header]
        for section_key, _ in self._section_order:
            content = sections.get(section_key, "")
            if content.strip():
                parts.append(content.strip())
                parts.append("\n\n---\n\n")
        return "".join(parts)

    def _save_markdown(self, content: str) -> Path:
        md_path = self.output_dir / "generated_task.md"
        md_path.write_text(content, encoding="utf-8")
        return md_path

    def _export_docx(self, markdown_content: str) -> Path:
        """将 Markdown 内容导出为 .docx 文件（使用 python-docx）。"""
        docx_path = self.output_dir / "generated_task.docx"
        try:
            from docx import Document
            from docx.shared import Pt

            doc = Document()
            for line in markdown_content.splitlines():
                stripped = line.strip()
                if stripped.startswith("# "):
                    doc.add_heading(stripped[2:], level=1)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("### "):
                    doc.add_heading(stripped[4:], level=3)
                elif stripped == "---":
                    doc.add_paragraph("─" * 40)
                elif stripped:
                    doc.add_paragraph(stripped)
            doc.save(str(docx_path))
        except Exception as e:
            logger.warning(f"docx export failed: {e}. Saving empty placeholder.")
            docx_path.write_bytes(b"")

        return docx_path


__all__ = ["TaskGenerator", "SECTION_ORDER_ZH", "SECTION_ORDER_EN"]
