# -*- coding: utf-8 -*-
"""
Project Creator 后端集成测试

不依赖真实 claude CLI——只测试 REST 端点和会话 CRUD。
使用 httpx.AsyncClient + ASGITransport 直接调用 FastAPI app。
"""
import io
import zipfile

import pytest
import pytest_asyncio
from httpx import AsyncClient, ASGITransport

from src.api.main import app


# ─── Fixtures ──────────────────────────────────────────────────────────────────


@pytest_asyncio.fixture
async def client():
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as c:
        yield c


@pytest.fixture
def minimal_docx_bytes() -> bytes:
    """使用 python-docx 创建最小合法 .docx 文件（含章节关键字）"""
    from docx import Document

    doc = Document()
    doc.add_heading("课程目标", level=1)
    doc.add_paragraph("测试内容")
    doc.add_heading("模块一", level=2)
    doc.add_paragraph("模块内容描述")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── 1. upload-reference ───────────────────────────────────────────────────────


@pytest.mark.anyio
async def test_upload_reference_valid_docx(client, minimal_docx_bytes):
    """上传合法 .docx → 200 + structure 字段"""
    response = await client.post(
        "/api/v1/project/upload-reference",
        files={
            "files": (
                "test_ref.docx",
                minimal_docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 200
    data = response.json()
    assert "reference_id" in data
    assert "structure" in data
    assert data["original_filename"] == "test_ref.docx"
    assert "sections" in data["structure"]


@pytest.mark.anyio
async def test_upload_reference_invalid_type(client):
    """上传非 docx/pdf 文件 → 400"""
    response = await client.post(
        "/api/v1/project/upload-reference",
        files={"files": ("test.txt", b"hello world", "text/plain")},
    )
    assert response.status_code == 400
    assert "Unsupported file type" in response.json()["detail"]


@pytest.mark.anyio
async def test_upload_reference_no_file(client):
    """不提供文件 → 422（FastAPI 参数校验）"""
    response = await client.post("/api/v1/project/upload-reference")
    assert response.status_code == 422


# ─── 2. sessions CRUD ──────────────────────────────────────────────────────────


@pytest.mark.anyio
async def test_list_sessions_returns_list(client):
    """GET /project/sessions → 200 + sessions 列表"""
    response = await client.get("/api/v1/project/sessions")
    assert response.status_code == 200
    data = response.json()
    assert "sessions" in data
    assert "total" in data
    assert isinstance(data["sessions"], list)


@pytest.mark.anyio
async def test_get_session_not_found(client):
    """不存在的 session_id → 404"""
    response = await client.get("/api/v1/project/sessions/nonexistent-id-99999")
    assert response.status_code == 404


@pytest.mark.anyio
async def test_delete_session_not_found(client):
    """删除不存在的 session → 404"""
    response = await client.delete("/api/v1/project/sessions/nonexistent-id-99999")
    assert response.status_code == 404


@pytest.mark.anyio
async def test_session_full_crud(client):
    """创建 → GET → DELETE → GET(404) 完整流程"""
    from src.agents.project import get_project_session_manager

    mgr = get_project_session_manager()
    sid = mgr.create_session(
        theme="集成测试主题",
        kb_name=None,
        reference_structure={},
    )
    try:
        # GET existing
        r = await client.get(f"/api/v1/project/sessions/{sid}")
        assert r.status_code == 200
        assert r.json()["theme"] == "集成测试主题"

        # DELETE
        r = await client.delete(f"/api/v1/project/sessions/{sid}")
        assert r.status_code == 200
        assert r.json()["deleted"] == sid

        # GET after delete → 404
        r = await client.get(f"/api/v1/project/sessions/{sid}")
        assert r.status_code == 404
    finally:
        # safety cleanup (no-op if already deleted)
        mgr.delete_session(sid)


@pytest.mark.anyio
async def test_list_sessions_limit(client):
    """limit 参数被接受，返回不超过 limit 条"""
    response = await client.get("/api/v1/project/sessions?limit=2")
    assert response.status_code == 200
    data = response.json()
    assert len(data["sessions"]) <= 2


# ─── 3. download-task ──────────────────────────────────────────────────────────


@pytest.mark.anyio
async def test_download_task_session_not_found(client):
    """session 不存在 → 404"""
    response = await client.get("/api/v1/project/nonexistent-id/download-task")
    assert response.status_code == 404


@pytest.mark.anyio
async def test_download_task_md_success(client, tmp_path):
    """session 存在且 md 文件存在 → 返回文件"""
    from src.agents.project import get_project_session_manager

    mgr = get_project_session_manager()
    sid = mgr.create_session(theme="下载测试", kb_name=None, reference_structure={})
    try:
        md_file = tmp_path / "task.md"
        md_file.write_text("# 测试任务书\n\n内容", encoding="utf-8")
        mgr.update_session(sid, task_md_path=str(md_file))

        response = await client.get(f"/api/v1/project/{sid}/download-task?format=md")
        assert response.status_code == 200
        assert "content-disposition" in response.headers
    finally:
        mgr.delete_session(sid)


# ─── 4. download-repo ──────────────────────────────────────────────────────────


@pytest.mark.anyio
async def test_download_repo_session_not_found(client):
    """session 不存在 → 404"""
    response = await client.get("/api/v1/project/nonexistent-id/download-repo")
    assert response.status_code == 404


@pytest.mark.anyio
async def test_download_repo_no_repo_dir(client):
    """session 存在但无 repo 目录 → 404"""
    from src.agents.project import get_project_session_manager

    mgr = get_project_session_manager()
    sid = mgr.create_session(theme="无仓库测试", kb_name=None, reference_structure={})
    try:
        response = await client.get(f"/api/v1/project/{sid}/download-repo")
        assert response.status_code == 404
    finally:
        mgr.delete_session(sid)


@pytest.mark.anyio
async def test_download_repo_returns_zip(client):
    """repo 目录存在 → 返回 application/zip，内含文件"""
    import shutil
    from src.agents.project import get_project_session_manager
    from src.api.routers.project import PROJECTS_DIR

    mgr = get_project_session_manager()
    sid = mgr.create_session(theme="仓库下载测试", kb_name=None, reference_structure={})

    repo_dir = PROJECTS_DIR / sid / "repo"
    repo_dir.mkdir(parents=True, exist_ok=True)
    (repo_dir / "README.md").write_text("# Test Repo", encoding="utf-8")
    (repo_dir / "main.py").write_text("print('hello')", encoding="utf-8")

    try:
        response = await client.get(f"/api/v1/project/{sid}/download-repo")
        assert response.status_code == 200
        assert response.headers["content-type"] == "application/zip"

        buf = io.BytesIO(response.content)
        with zipfile.ZipFile(buf) as zf:
            names = zf.namelist()
        assert "README.md" in names
        assert "main.py" in names
        # .git not included
        assert not any(".git" in n for n in names)
    finally:
        shutil.rmtree(PROJECTS_DIR / sid, ignore_errors=True)
        mgr.delete_session(sid)
